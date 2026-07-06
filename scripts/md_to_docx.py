"""将 full_paper.md 转换为高质量中文学术格式 DOCX"""
import re, os, sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement

# LaTeX 希腊字母 → Unicode 映射
LATEX_TO_UNICODE = {
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\epsilon': 'ε', r'\varepsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η',
    r'\theta': 'θ', r'\vartheta': 'ϑ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ',
    r'\pi': 'π', r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ',
    r'\upsilon': 'υ', r'\phi': 'φ', r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
    r'\Xi': 'Ξ', r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Upsilon': 'Υ',
    r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
    r'\infty': '∞', r'\approx': '≈', r'\neq': '≠', r'\leq': '≤',
    r'\geq': '≥', r'\pm': '±', r'\times': '×', r'\cdot': '·',
    r'\rightarrow': '→', r'\Rightarrow': '⇒', r'\leftarrow': '←',
}

def latex_to_word_runs(paragraph, text):
    """将含 $...$ 内联公式的文本渲染为 Word runs，处理希腊字母和上下标"""
    # 先处理所有 $...$ 段
    segments = re.split(r'(\$[^$]+\$)', text)
    for seg in segments:
        if seg.startswith('$') and seg.endswith('$'):
            formula = seg[1:-1]  # 去掉 $ 包裹
            _render_formula(paragraph, formula)
        elif seg:
            # 处理 **加粗**
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', seg)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = paragraph.add_run(bp[2:-2])
                    run.bold = True
                    run.font.size = Pt(12)
                    set_cn_font(run, '宋体')
                elif bp:
                    run = paragraph.add_run(bp)
                    run.font.size = Pt(12)
                    set_cn_font(run, '宋体')

def _render_formula(paragraph, formula):
    """将单个 LaTeX 公式片段渲染为 Word runs（希腊字母 + 上下标 + 斜体变量）"""
    # 替换 LaTeX 命令为 Unicode
    rendered = formula
    for latex, uni in sorted(LATEX_TO_UNICODE.items(), key=lambda x: -len(x[0])):
        rendered = rendered.replace(latex, uni)

    # 处理 _{...} 和 _单字符 下标、^{...} 和 ^单字符 上标
    pattern = re.compile(r'([_\^])(?:\{([^}]+)\}|(\w))')
    pos = 0
    last_end = 0

    for m in pattern.finditer(rendered):
        # 输出前面的普通文本
        plain = rendered[last_end:m.start()]
        if plain:
            run = paragraph.add_run(plain)
            run.font.size = Pt(12)
            run.italic = True
            set_cn_font(run, 'Cambria Math')

        # 输出上/下标
        sub_sup_text = m.group(2) or m.group(3)
        run = paragraph.add_run(sub_sup_text)
        run.font.size = Pt(9)
        if m.group(1) == '_':
            run.font.subscript = True
        else:
            run.font.superscript = True
        set_cn_font(run, 'Cambria Math')

        last_end = m.end()

    # 剩余文本
    remaining = rendered[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.size = Pt(12)
        run.italic = True
        set_cn_font(run, 'Cambria Math')

def _render_block_formula(doc, formula):
    """渲染块级公式 $$...$$"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.5, space_before=Pt(6), space_after=Pt(6))

    rendered = formula
    for latex, uni in sorted(LATEX_TO_UNICODE.items(), key=lambda x: -len(x[0])):
        rendered = rendered.replace(latex, uni)

    # 处理上下标
    pattern = re.compile(r'([_\^])(?:\{([^}]+)\}|(\w))')
    pos = 0
    last_end = 0

    for m in pattern.finditer(rendered):
        plain = rendered[last_end:m.start()]
        if plain:
            run = p.add_run(plain)
            run.font.size = Pt(11)
            run.italic = True
            set_cn_font(run, 'Cambria Math')

        sub_sup_text = m.group(2) or m.group(3)
        run = p.add_run(sub_sup_text)
        run.font.size = Pt(9)
        if m.group(1) == '_':
            run.font.subscript = True
        else:
            run.font.superscript = True
        set_cn_font(run, 'Cambria Math')

        last_end = m.end()

    remaining = rendered[last_end:]
    if remaining:
        run = p.add_run(remaining)
        run.font.size = Pt(11)
        run.italic = True
        set_cn_font(run, 'Cambria Math')

    return p


def set_cn_font(run, font_name='宋体', font_name_east='宋体'):
    """设置中文字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name_east)
    rFonts.set(qn('w:hAnsi'), font_name_east)

def set_paragraph_spacing(paragraph, line_spacing=1.5, space_after=Pt(0), space_before=Pt(0)):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before

def add_body_paragraph(doc, text, indent_first_line=True):
    """添加正文段落"""
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5)
    if indent_first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    latex_to_word_runs(p, text)
    return p

def add_heading_cn(doc, text, level=1):
    """添加中文标题 (黑体加粗)"""
    text = text.strip()
    h = doc.add_paragraph()
    set_paragraph_spacing(h, line_spacing=1.5, space_before=Pt(12), space_after=Pt(6))
    run = h.add_run(text)
    run.bold = True
    set_cn_font(run, '黑体', '黑体')
    sizes = {0: 18, 1: 16, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 0:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def _write_table(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    for ri, row_cells in enumerate(rows):
        for ci in range(num_cols):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].clear()
            cell_text = row_cells[ci] if ci < len(row_cells) else ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)
            if ri == 0:
                run.bold = True
                set_cn_font(run, '黑体')
            else:
                set_cn_font(run, '宋体')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def md_to_docx(md_path, docx_path):
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    table_data = []
    in_table = False
    in_formula = False
    formula_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            if in_table and table_data:
                _write_table(doc, table_data)
                table_data = []
                in_table = False
            i += 1
            continue

        if line.strip() == '---':
            i += 1
            continue

        # $$ 块级公式
        if line.strip() == '$$' and not in_formula:
            in_formula = True
            formula_lines = []
            i += 1
            continue
        if line.strip() == '$$' and in_formula:
            in_formula = False
            formula = ' '.join(formula_lines)
            _render_block_formula(doc, formula)
            i += 1
            continue
        if in_formula:
            formula_lines.append(line.strip())
            i += 1
            continue

        # 表格
        if line.strip().startswith('|'):
            in_table = True
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(re.match(r'^[-:—]+$', c) for c in cells):
                i += 1
                continue
            table_data.append(cells)
            i += 1
            continue
        elif in_table and table_data:
            _write_table(doc, table_data)
            table_data = []
            in_table = False

        # 标题层级
        if line.startswith('# ') and not line.startswith('## '):
            add_heading_cn(doc, line[2:].strip(), level=0)
            i += 1; continue
        if line.startswith('## ') and not line.startswith('### '):
            add_heading_cn(doc, line[3:].strip(), level=1)
            i += 1; continue
        if line.startswith('### ') and not line.startswith('#### '):
            add_heading_cn(doc, line[4:].strip(), level=2)
            i += 1; continue
        if line.startswith('#### '):
            add_heading_cn(doc, line[5:].strip(), level=3)
            i += 1; continue

        # 参考文献条目
        if re.match(r'^\[\d+\]', line):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.5)
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-1.5)
            run = p.add_run(line.strip())
            run.font.size = Pt(10.5)
            set_cn_font(run, '宋体')
            i += 1; continue

        # 假说行
        if re.match(r'^\*\*H\d\*\*', line):
            add_body_paragraph(doc, line, indent_first_line=False)
            i += 1; continue

        # 表注
        if line.strip().startswith('注：'):
            p = add_body_paragraph(doc, line, indent_first_line=False)
            for run in p.runs:
                run.font.size = Pt(9)
            i += 1; continue

        # 编号列表项
        m = re.match(r'^\*\*（(\d+)）\*\*(.*)', line)
        if m:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.5)
            p.paragraph_format.left_indent = Cm(1)
            run1 = p.add_run(f'（{m.group(1)}）')
            run1.bold = True; run1.font.size = Pt(12); set_cn_font(run1, '黑体')
            if m.group(2).strip():
                run2 = p.add_run(m.group(2))
                run2.font.size = Pt(12); set_cn_font(run2, '宋体')
            i += 1; continue

        # 普通段落
        add_body_paragraph(doc, line)
        i += 1

    if table_data:
        _write_table(doc, table_data)

    doc.save(docx_path)
    size = os.path.getsize(docx_path)
    print(f'[OK] DOCX: {docx_path} ({size//1024} KB)')

if __name__ == '__main__':
    md_path = sys.argv[1] if len(sys.argv) > 1 else 'workspace/writing/full_paper.md'
    docx_path = sys.argv[2] if len(sys.argv) > 2 else 'workspace/writing/full_paper.docx'
    md_to_docx(md_path, docx_path)
